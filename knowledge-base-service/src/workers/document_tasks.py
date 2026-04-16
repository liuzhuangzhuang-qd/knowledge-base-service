import os

from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from src.core.db import SessionLocal
from src.models import ChunkEmbedding, Document, DocumentChunk
from src.services.chunking import split_text
from src.services.qwen_client import embed_texts


def _normalize_text(text: str) -> str:
    """
    Normalize extracted document text for chunking/retrieval quality.
    """
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    cleaned_lines: list[str] = []
    prev_blank = False

    for line in lines:
        normalized = " ".join(line.strip().split())
        if not normalized:
            if not prev_blank:
                cleaned_lines.append("")
            prev_blank = True
            continue
        prev_blank = False
        cleaned_lines.append(normalized)

    return "\n".join(cleaned_lines).strip()


def _load_raw_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext in {".txt", ".md"}:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == ".docx":
        doc = DocxDocument(file_path)
        paragraph_lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        table_lines: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))
        return "\n".join(paragraph_lines + table_lines)
    raise ValueError(f"Unsupported file extension: {ext}")


def process_document(doc_id: int) -> None:
    db: Session = SessionLocal()
    try:
        document = db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            return

        document.status = "parsing"
        db.commit()

        raw_text = _load_raw_text(document.file_path)

        raw_length = len(raw_text)
        normalized_text = _normalize_text(raw_text)
        normalized_length = len(normalized_text)
        dropped_ratio = 0.0
        if raw_length > 0:
            dropped_ratio = round((raw_length - normalized_length) / raw_length, 4)

        if not normalized_text:
            raise ValueError("Document contains no extractable text")

        document.status = "chunking"
        db.commit()
        chunks = split_text(normalized_text)

        db.query(ChunkEmbedding).filter(
            ChunkEmbedding.chunk_id.in_(
                db.query(DocumentChunk.id).filter(DocumentChunk.doc_id == doc_id)
            )
        ).delete(synchronize_session=False)
        db.query(DocumentChunk).filter(DocumentChunk.doc_id == doc_id).delete(
            synchronize_session=False
        )
        db.commit()

        created_chunks: list[DocumentChunk] = []
        for i, content in enumerate(chunks):
            c = DocumentChunk(doc_id=doc_id, chunk_index=i, content=content)
            db.add(c)
            created_chunks.append(c)
        db.commit()

        document.status = "embedding"
        db.commit()
        vectors = embed_texts([c.content for c in created_chunks]) if created_chunks else []
        for c, vec in zip(created_chunks, vectors):
            db.add(ChunkEmbedding(chunk_id=c.id, vector=vec))
        db.commit()

        document.status = "ready"
        document.metadata_json = {
            "chunk_count": len(created_chunks),
            "raw_length": raw_length,
            "normalized_length": normalized_length,
            "dropped_ratio": dropped_ratio,
        }
        db.commit()
    except Exception as exc:
        document = db.query(Document).filter(Document.id == doc_id).first()
        if document:
            document.status = "failed"
            document.metadata_json = {"error": str(exc)}
            db.commit()
    finally:
        db.close()
